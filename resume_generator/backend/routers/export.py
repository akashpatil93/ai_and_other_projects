import io
import re
from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from services import session_store

router = APIRouter()

# ─── Content Cleaning ─────────────────────────────────────────────────────────

_RESUME_SECTIONS = {
    "SUMMARY", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EDUCATION",
    "SKILLS", "CERTIFICATIONS", "PROJECTS", "ACHIEVEMENTS", "EXPERIENCE",
    "OBJECTIVE", "PROFILE", "INTERNSHIP EXPERIENCE", "LEADERSHIP",
    "LEADERSHIP & RECOGNITION", "EDUCATIONAL QUALIFICATIONS",
    "AWARDS", "PUBLICATIONS", "VOLUNTEER", "HONORS", "ADDITIONAL",
}

_CONVERSATIONAL_STARTERS = (
    "let me know", "please let", "would you like", "feel free",
    "once you", "if you'd", "if you ", "i can also", "i can generate",
    "you can now", "head to", "reach out", "happy to", "don't hesitate",
    "you're satisfied", "i'm happy", "let me generate", "let me now",
    "here is", "here's", "below is", "i've tailored", "i have tailored",
    "i'll now", "i will now", "minor gap", "match analysis",
    "akash is a", "he has ", "the jd ",
)

_SEPARATOR_RE = re.compile(r'^[\s■▪•\-\*_=|]+$')


def _strip_markdown(text: str) -> str:
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*([^*\n]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*\n]+)\*', r'\1', text)
    return text


def _is_section_heading(raw: str) -> bool:
    s = re.sub(r'[\*#■▪]+', '', raw).strip()
    if not s or len(s) > 65:
        return False
    return s.upper() in _RESUME_SECTIONS


def _extract_resume_content(text: str) -> str:
    lines = text.split("\n")

    # Find where the actual resume starts
    start_idx = 0
    for i, raw_line in enumerate(lines):
        clean = re.sub(r'[\*#■▪]+', '', raw_line).strip()
        is_section = _is_section_heading(clean)
        is_contact = "|" in clean and ("@" in clean or re.search(r'\+\d', clean))
        if is_section or is_contact:
            # Walk back to capture candidate name (skip separator-only lines)
            start_idx = i
            for j in range(i - 1, max(-1, i - 6), -1):
                stripped = re.sub(r'[\s■▪•\-\*_=|]', '', lines[j])
                if stripped:
                    start_idx = j
                    break
            break

    # Strip trailing conversational lines
    end_idx = len(lines)
    for i in range(len(lines) - 1, start_idx - 1, -1):
        line = lines[i].strip().lower()
        if not line or _SEPARATOR_RE.match(lines[i]):
            continue
        if any(line.startswith(s) for s in _CONVERSATIONAL_STARTERS):
            end_idx = i
        else:
            break

    result = lines[start_idx:end_idx]
    while result and not result[0].strip():
        result.pop(0)
    while result and not result[-1].strip():
        result.pop()
    return "\n".join(result)


def _clean_content(text: str) -> str:
    return _strip_markdown(_extract_resume_content(text))


# ─── Structural Parser ────────────────────────────────────────────────────────

_DATE_PAT = re.compile(
    r'(Present|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})',
    re.I,
)


def _parse_job_line(line: str):
    """Return (company, role, date) if line looks like a job entry, else None."""
    s = line.strip()
    if "|" not in s:
        return None
    parts = [p.strip() for p in s.split("|")]
    if len(parts) >= 2 and _DATE_PAT.search(parts[-1]):
        company = parts[0]
        date = parts[-1]
        role = " | ".join(parts[1:-1])
        return company, role, date
    return None


def _looks_like_job_line(line: str) -> bool:
    """Company  Role  Date pattern without pipe separator."""
    s = line.strip()
    if not s or s[0] in "•▪-*■":
        return False
    if _is_section_heading(s):
        return False
    return bool(_DATE_PAT.search(s)) and len(s) < 120


def _parse_resume(content: str):
    """
    Returns (name, contact_parts, items).
    Items: ("section"|"job"|"subhdr"|"bullet"|"body", *values)
    """
    lines = [ln.rstrip() for ln in content.split("\n")]

    # Remove separator-only lines everywhere
    lines = [ln for ln in lines if not _SEPARATOR_RE.match(ln)]

    # Name: first non-empty line
    name, pos = "", 0
    for i, ln in enumerate(lines):
        if ln.strip():
            name = ln.strip()
            pos = i + 1
            break

    # Contact lines immediately after name
    contact_parts = []
    for i in range(pos, min(pos + 5, len(lines))):
        ln = lines[i].strip()
        if not ln:
            continue
        is_contact = (
            ("|" in ln and ("@" in ln or re.search(r'\+\d', ln) or "linkedin" in ln.lower()))
            or bool(re.search(r'^\+\d', ln))
            or ("@" in ln and not _is_section_heading(ln) and len(ln) < 80)
            or "linkedin.com" in ln.lower()
        )
        if is_contact:
            parts = [p.strip() for p in ln.split("|") if p.strip()]
            contact_parts.extend(parts if parts else [ln])
            pos = i + 1
        else:
            break

    items = []
    in_job = False

    for raw in lines[pos:]:
        ln = raw.strip()
        if not ln:
            continue

        if _is_section_heading(ln):
            items.append(("section", re.sub(r'[\*#■▪]+', '', ln).strip()))
            in_job = False
            continue

        # Pipe-separated job entry
        job = _parse_job_line(ln)
        if job:
            items.append(("job",) + job)
            in_job = True
            continue

        # Bullet
        if ln[0] in "•▪-*■" and len(ln) > 2:
            clean = re.sub(r'^[•▪\-\*■\s]+', '', ln).strip()
            if clean:
                items.append(("bullet", clean))
            continue

        # Plain text
        if in_job:
            items.append(("subhdr", ln))
        else:
            items.append(("body", ln))

    return name, contact_parts, items


# ─── Endpoint ─────────────────────────────────────────────────────────────────

class ExportRequest(BaseModel):
    session_id: str
    content_type: str   # "resume" | "cover_letter"
    format: str         # "txt" | "pdf" | "docx"


@router.post("/download")
async def export_document(request: ExportRequest):
    session = session_store.get_session(request.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found.")

    if request.content_type == "resume":
        content = session.get("generated_resume", "")
        base_name = "resume"
    elif request.content_type == "cover_letter":
        content = session.get("generated_cover_letter", "")
        base_name = "cover_letter"
    else:
        raise HTTPException(status_code=400, detail="Invalid content_type.")

    if not content:
        raise HTTPException(status_code=400, detail="No content to export yet.")

    content = _clean_content(content)
    fmt = request.format.lower()

    if fmt == "txt":
        return Response(
            content=content.encode("utf-8"),
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename={base_name}.txt"},
        )
    elif fmt == "pdf":
        gen = _generate_resume_pdf if request.content_type == "resume" else _generate_simple_pdf
        return Response(
            content=gen(content),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={base_name}.pdf"},
        )
    elif fmt == "docx":
        gen = _generate_resume_docx if request.content_type == "resume" else _generate_simple_docx
        return Response(
            content=gen(content),
            media_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            headers={"Content-Disposition": f"attachment; filename={base_name}.docx"},
        )
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use txt, pdf, or docx.")


# ─── PDF — Resume ─────────────────────────────────────────────────────────────

def _generate_resume_pdf(content: str) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    )
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT

    name, contact_parts, items = _parse_resume(content)

    LM = RM = 0.65 * inch
    TM = BM = 0.6 * inch
    PAGE_W = letter[0] - LM - RM

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=LM, rightMargin=RM,
        topMargin=TM, bottomMargin=BM,
    )

    BLACK  = colors.HexColor("#111111")
    DGRAY  = colors.HexColor("#333333")
    MGRAY  = colors.HexColor("#555555")
    RULE   = colors.HexColor("#CCCCCC")

    def sty(name, **kw):
        base = dict(fontName="Helvetica", fontSize=9.5, textColor=DGRAY,
                    leading=13.5, spaceAfter=0, spaceBefore=0, alignment=TA_LEFT)
        base.update(kw)
        return ParagraphStyle(name, **base)

    name_sty    = sty("N",  fontSize=20, fontName="Helvetica-Bold",
                       textColor=BLACK, leading=24, spaceAfter=2)
    contact_sty = sty("C",  fontSize=9, textColor=MGRAY, alignment=TA_LEFT)
    sec_sty     = sty("S",  fontSize=10.5, fontName="Helvetica-Bold",
                       textColor=BLACK, spaceBefore=10, spaceAfter=2)
    company_sty = sty("Co", fontSize=10, fontName="Helvetica-Bold", textColor=BLACK)
    role_sty    = sty("R",  fontSize=9.5, textColor=DGRAY, alignment=TA_CENTER)
    date_sty    = sty("D",  fontSize=9.5, textColor=DGRAY, alignment=TA_RIGHT)
    subhdr_sty  = sty("Sh", fontSize=9.5, fontName="Helvetica-BoldOblique",
                       textColor=DGRAY, spaceBefore=3, spaceAfter=1)
    bullet_sty  = sty("B",  fontSize=9.5, leftIndent=12, firstLineIndent=0,
                       spaceAfter=1, spaceBefore=1)
    body_sty    = sty("Bo", fontSize=9.5, spaceAfter=2)

    story = []

    # ── Name ──
    story.append(Paragraph(name, name_sty))

    # ── Contact line ──
    if contact_parts:
        story.append(Paragraph("  |  ".join(contact_parts), contact_sty))

    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=0.8, color=BLACK, spaceAfter=4))

    # ── Body items ──
    for item in items:
        kind = item[0]

        if kind == "section":
            story.append(Spacer(1, 4))
            story.append(Paragraph(item[1].upper(), sec_sty))
            story.append(HRFlowable(width="100%", thickness=0.5, color=RULE, spaceAfter=3))

        elif kind == "job":
            company, role, date = item[1], item[2], item[3]
            col_w = [PAGE_W * 0.40, PAGE_W * 0.33, PAGE_W * 0.27]
            t = Table(
                [[Paragraph(company, company_sty),
                  Paragraph(role or "", role_sty),
                  Paragraph(date, date_sty)]],
                colWidths=col_w,
            )
            t.setStyle(TableStyle([
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING",   (0, 0), (-1, -1), 0),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
                ("TOPPADDING",    (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(t)

        elif kind == "subhdr":
            story.append(Paragraph(item[1], subhdr_sty))

        elif kind == "bullet":
            story.append(Paragraph(f"\u2022  {item[1]}", bullet_sty))

        elif kind == "body":
            story.append(Paragraph(item[1], body_sty))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ─── PDF — Cover Letter ───────────────────────────────────────────────────────

def _generate_simple_pdf(content: str) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch,
    )
    body = ParagraphStyle(
        "CL", fontName="Helvetica", fontSize=10.5,
        textColor=colors.HexColor("#111111"), leading=16,
    )
    story = []
    for line in content.split("\n"):
        if line.strip():
            story.append(Paragraph(line.strip(), body))
        else:
            story.append(Spacer(1, 10))
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ─── DOCX — Resume ────────────────────────────────────────────────────────────

def _generate_resume_docx(content: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    name, contact_parts, items = _parse_resume(content)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.65)
        sec.bottom_margin = Inches(0.65)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    BLACK = RGBColor(0x11, 0x11, 0x11)
    GRAY  = RGBColor(0x55, 0x55, 0x55)

    def _no_space(p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    def _add_bottom_border(p, color="AAAAAA", sz="4"):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    sz)
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _remove_cell_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            tcBorders.append(el)
        tcPr.append(tcBorders)

    # ── Name ──
    p = doc.add_paragraph()
    _no_space(p)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLACK

    # ── Contact ──
    if contact_parts:
        p = doc.add_paragraph()
        _no_space(p)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("  |  ".join(contact_parts))
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY

    # Thick rule under header
    p = doc.add_paragraph()
    _no_space(p)
    _add_bottom_border(p, color="111111", sz="8")

    # ── Body items ──
    for item in items:
        kind = item[0]

        if kind == "section":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(item[1].upper())
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = BLACK
            _add_bottom_border(p)

        elif kind == "job":
            company, role, date = item[1], item[2], item[3]
            t = doc.add_table(rows=1, cols=3)
            for cell in t.rows[0].cells:
                _remove_cell_borders(cell)

            c0 = t.rows[0].cells[0]
            p0 = c0.paragraphs[0]
            _no_space(p0)
            p0.paragraph_format.space_before = Pt(5)
            r0 = p0.add_run(company)
            r0.bold = True
            r0.font.size = Pt(10)
            r0.font.color.rgb = BLACK

            c1 = t.rows[0].cells[1]
            p1 = c1.paragraphs[0]
            _no_space(p1)
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(role or "")
            r1.font.size = Pt(9.5)

            c2 = t.rows[0].cells[2]
            p2 = c2.paragraphs[0]
            _no_space(p2)
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r2 = p2.add_run(date)
            r2.font.size = Pt(9.5)

        elif kind == "subhdr":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.1)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(item[1])
            r.bold = True
            r.italic = True
            r.font.size = Pt(9.5)

        elif kind == "bullet":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.space_before      = Pt(1)
            p.paragraph_format.space_after       = Pt(1)
            r = p.add_run(f"\u2022  {item[1]}")
            r.font.size = Pt(9.5)

        elif kind == "body":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(item[1])
            r.font.size = Pt(9.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ─── DOCX — Cover Letter ──────────────────────────────────────────────────────

def _generate_simple_docx(content: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    for line in content.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0 if not line.strip() else 6)
        if line.strip():
            r = p.add_run(line.strip())
            r.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
